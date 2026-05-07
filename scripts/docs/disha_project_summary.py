import sys, os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Colour palette ────────────────────────────────────────────────────────────
DEEP_BLUE = RGBColor(0x1A, 0x36, 0x5D)
NAVY      = RGBColor(0x2D, 0x5A, 0x87)
CORAL     = RGBColor(0xFF, 0x6B, 0x6B)
SUCCESS   = RGBColor(0x10, 0xB9, 0x81)
TEXT_MED  = RGBColor(0x47, 0x55, 0x69)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

def set_color(run, rgb): run.font.color.rgb = rgb
def bold(run): run.bold = True

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = DEEP_BLUE if level == 1 else NAVY
    run.font.name = "Calibri"
    return p

def add_para(text, color=None, size=10.5, bold_=False, italic_=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold_
    run.italic = italic_
    if color: run.font.color.rgb = color
    return p

def add_bullet(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            text, is_bold = item
        else:
            text, is_bold = item, False
        run = p.add_run(text)
        run.bold = is_bold
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = TEXT_MED

def add_callout(text, label="IMPORTANT", bg=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r1 = p.add_run(f"  {label}:  ")
    r1.bold = True
    r1.font.color.rgb = WHITE
    r1.font.name = "Calibri"
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = DEEP_BLUE
    return p

def divider():
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
        run.font.size = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Celebrate Festival Inc — Project Delivery Summary")
r.bold = True; r.font.size = Pt(20); r.font.name = "Calibri"
r.font.color.rgb = DEEP_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("Prepared for: Disha  |  Hiraya Digital  |  May 2026")
r2.font.size = Pt(10); r2.font.name = "Calibri"; r2.font.color.rgb = TEXT_MED; r2.italic = True

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PURPOSE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Purpose of This Document")
add_para(
    "This document provides a complete, honest account of what has been delivered on the "
    "Celebrate Festival Inc project — what was in the original scope, what was added beyond it, "
    "and where client-side actions are still blocking certain features. "
    "It is intended to help you, as account lead, communicate the full value of this engagement "
    "to the client and to establish a clear boundary between Phase 1 (complete) and any future work.",
    color=TEXT_MED
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ORIGINAL SCOPE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Original Scope of Work (PDF Sent: May 19)")
add_para(
    "The original service breakdown document defined a standard Shopify website redesign. "
    "Core deliverables included:",
    color=TEXT_MED
)
add_bullet([
    "New homepage design with featured products",
    "Redesigned product pages with improved galleries",
    "Mobile-friendly responsive design",
    "Standard navigation structure and breadcrumb navigation",
    "Standard Shopify search and basic product filtering",
    "Stock availability indicators",
    "Standard checkout process",
    "Basic email marketing setup using Shopify Email",
    "Basic cross-sell sections and related products",
    "Customer account area with standard features",
    "Standard SEO optimisation",
    "Basic testing, QA, site launch and post-launch support",
])
add_para(
    "This is the work a typical Shopify theme developer completes. "
    "What was actually delivered is in a completely different category.",
    color=CORAL, bold_=True, size=10.5
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. WHAT WAS ACTUALLY BUILT
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. What Was Actually Built — Beyond Any Standard Scope")

add_heading("3.1  Custom Theme — Built From Scratch", level=2)
add_para(
    "This is not a redesigned theme. This is a custom Shopify theme built from the ground up. "
    "A proper design process was followed: templates were designed, reviewed, and approved by the client. "
    "After approval and completion, the client requested multiple full redesign phases to match "
    "competitor websites — each phase requiring complete rework of already-finished templates.",
    color=TEXT_MED
)

add_heading("3.2  Three-Level Custom Collection Hierarchy (L1 / L2 / L3)", level=2)
add_para(
    "Three completely different collection page templates were built — each with its own layout, "
    "card design, image logic, hero banner, and navigation behaviour:",
    color=TEXT_MED
)
add_bullet([
    ("L1 — Main Category Page", True) ,
    "Hero banner, 4-column category grid, icon row, brands row, top products carousel",
    ("L2 — Subcategory Page", True),
    "3-column category cards, 8-column icon grid, AJAX product loading, sub-navigation",
    ("L3 — Product Listing Page", True),
    "Filter sidebar (260px), 4-column product grid, power type selector, toolbar with sort/view toggle",
])
add_para(
    "None of this exists natively in Shopify. Every template is custom Liquid + CSS + JavaScript.",
    color=CORAL, bold_=True
)

add_heading("3.3  Wholesale Pricing Engine — WSH BDR Integration", level=2)
add_para(
    "The store uses WSH (Wholesale Helper) on the Global/BDR plan. "
    "On this plan, Shopify's Liquid templating system cannot access live wholesale prices — "
    "the Liquid metafields are not refreshed at render time. "
    "This was confirmed directly by WSH support (Abdullah, Feb 10, 2026). "
    "This is not a bug — it is a documented platform limitation.",
    color=TEXT_MED
)
add_para("To work around this, a complete client-side pricing engine was built:", color=TEXT_MED)
add_bullet([
    "Intercepts WSH's own API calls to capture the live authentication token",
    "Calls the BDR (Backend Data Retrieval) API client-side on every page load",
    "Evaluates user type: CPH member, ROU member, general wholesale, or non-member",
    "Evaluates product vendor: CF Inc own-brand vs third-party",
    "Based on this matrix, displays the correct pricing UI for each user on each product",
    "Implemented across 6 page types: L3 collection, L2 hub, Search, Single Product Page, Cart, Homepage",
])
add_para(
    "This is not theme work. This is a custom JavaScript pricing engine integrated into a third-party "
    "wholesale app. It would be quoted as a standalone integration project by any agency.",
    color=CORAL, bold_=True
)

add_heading("3.4  Product Card — Multi-User Pricing Logic", level=2)
add_para(
    "Every product card across the entire site evaluates the following decision tree in real time:",
    color=TEXT_MED
)
add_bullet([
    "Is the user logged in?",
    "If yes — are they a CPH member, ROU member, or another wholesale tag?",
    "Is the product a CF Inc own-brand or a third-party vendor product?",
    "Does this variant have a WSH wholesale price?",
    "Based on all of the above: show Member Price with Was/Save, OR blurred price with Login CTA, OR single price with Call for Details CTA",
])
add_para(
    "This logic runs on search results, collection pages, the homepage, the cart, and the SPP — "
    "consistently, correctly, for every user type.",
    color=TEXT_MED
)

add_heading("3.5  SEO, AEO & Structured Data", level=2)
add_bullet([
    "Full competitor analysis conducted",
    ("Meta data added to 233 collections", True),
    "Structured data (schema.org) implemented across product, collection, and content pages",
    "Full AEO (Answer Engine Optimisation) based on 2026 standards — optimised for AI search engines, not just Google",
    "Short-term impression drop was expected and explained: replacing a live theme causes Google to re-crawl and re-index. This is normal and temporary.",
    "SEO is now running in the correct direction with a solid technical foundation",
])

add_heading("3.6  Contact Page — 4-Form System with Lead Capture", level=2)
add_bullet([
    "Four distinct contact forms: General Inquiry, Quote Request, Technical Support, Wholesale Partnership",
    "SweetAlert2 loading states and success/error handling",
    "Lead capture pipeline: every submission is sent to the chatbot backend and stored",
    "GA4 and LinkedIn conversion tracking on every form submission",
    "Equipment leasing modal (Radiance Capital integration)",
])

add_heading("3.7  AI Kitchen Consultant Chatbot (Proactive Pitch — Not Client-Requested)", level=2)
add_para(
    "This was not requested by the client. It was built proactively by Hiraya Digital as an "
    "SEO and conversion tool — an intent-based chatbot that understands commercial kitchen equipment, "
    "guides users to the right products, and captures leads. "
    "It is currently in beta and being pitched to the client as an additional service.",
    color=TEXT_MED
)
add_bullet([
    "Custom Node.js / Express backend hosted on Hiraya Digital's VPS",
    "MySQL database for conversation history, leads, and analytics",
    "Google Gemini 2.5 Flash AI model with product-aware prompting",
    "Admin dashboard with full session auth, rate limiting, CAPTCHA, and account lockout",
    "Weekly email reports with lead summaries",
    "Three operating modes: standard chat, product SKU ordering, pro consultation",
    "Real server costs — this is not free infrastructure",
])
add_para(
    "If the client adopts this, it is a separate product with a separate monthly charge. "
    "It is not included in the theme project.",
    color=CORAL, bold_=True
)

add_heading("3.8  The Bigger Picture — We Are Bending Shopify Beyond Its Limits", level=2)
add_para(
    "It is important that this is stated clearly:",
    color=TEXT_MED
)
add_para(
    "We have engineered Shopify in ways it was never designed to work.",
    color=DEEP_BLUE, bold_=True, size=12
)
add_para(
    "Shopify is a hosted e-commerce platform built for standard retail stores. "
    "The client's competitors — WebstaurantStore, KaTom, and similar — run fully custom-built platforms "
    "developed by dedicated engineering teams over many years. "
    "These platforms have no Shopify limitations because they are not on Shopify.",
    color=TEXT_MED
)
add_para(
    "The client wants WebstaurantStore-level functionality on a Shopify platform. "
    "We have made that possible — but only by working around Shopify's architecture at every layer: "
    "intercepting API calls, building client-side pricing engines, "
    "routing around Shopify's contact form limitations using our own server, "
    "and constructing a three-level collection hierarchy that does not exist natively in any Shopify theme. "
    "None of this is standard theme work. None of this was in the original scope.",
    color=TEXT_MED
)
add_para(
    "The client does not see this complexity because the end result looks like 'just a website.' "
    "That invisibility is a sign of good engineering — not a sign that the work was easy.",
    color=CORAL, bold_=True
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CLIENT-SIDE BLOCKING ISSUES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Client-Side Work Not Yet Done — Blocking Us")

add_heading("4.1  L2 Collection Pages — Collections Not Populated", level=2)
add_para(
    "A proper, non-AJAX L2 collection page template was built and is ready. "
    "It works correctly when collections are populated with products and sub-collections. "
    "The client has not populated these collections.",
    color=TEXT_MED
)
add_para(
    "Because the client's collections are empty, an AJAX-based fallback template was built "
    "as a temporary workaround to keep the pages functional. "
    "This has been communicated explicitly to the client:",
    color=TEXT_MED
)
add_bullet([
    ("The AJAX template is a workaround, not a permanent solution.", True),
    ("We will not fix bugs in the AJAX template.", True),
    "The correct fix is for the client to populate their collections — this is their responsibility.",
    "Once collections are populated, the proper L2 template will work without issues.",
])

add_heading("4.2  File Upload on Technical Support Form", level=2)
add_para(
    "The client has raised the file upload in the Technical Support contact form as a bug. "
    "This is not a bug. This is a Shopify platform limitation.",
    color=TEXT_MED
)
add_para(
    "On February 6, 2026, a detailed WhatsApp message (Ref: #178215540764679) was sent "
    "explaining that Shopify's built-in contact form system does not support file attachments, "
    "and presenting two options:",
    color=TEXT_MED
)
add_bullet([
    ("Option 1", True) ,
    "Build a custom file upload solution using our server (additional development, separate charge)",
    ("Option 2", True),
    "Use a paid Shopify app (~$10–15/month, e.g. UploadKit) with upload limits and third-party dependency",
])
add_para(
    "The client did not reply to this message. The feature then reappeared as a bug in the tracker. "
    "A solution has since been designed (see Section 5) but it requires client confirmation and is a chargeable item.",
    color=CORAL, bold_=True
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROPOSED PHASE 2 — CHATBOT + FILE UPLOAD + LEAD CRM
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Proposed Additional Features — Separate Quote Required")
add_para(
    "Three features have been designed that all depend on Hiraya Digital's VPS server infrastructure. "
    "These are not Shopify features. They are custom-built capabilities that extend what Shopify can do. "
    "All three would be bundled as part of the AI Chatbot App subscription:",
    color=TEXT_MED
)

add_heading("5.1  File Upload on Contact Forms", level=2)
add_bullet([
    "User selects a file in the Technical Support form",
    "JavaScript uploads the file directly to our server before form submission",
    "Server stores the file securely and returns a shareable link",
    "That link is injected into the Shopify contact form as a text field",
    "Shopify sends the email with the file link — no platform limitation",
    "Clean, no design change, no third-party subscription required",
    ("This requires our VPS server and is chargeable.", True),
])

add_heading("5.2  Contact Form Lead Capture → Shopify CRM", level=2)
add_bullet([
    "Every contact form submission automatically creates or updates a customer record in Shopify",
    "Tags applied based on form type (e.g. 'lead-technical-support', 'lead-wholesale')",
    "Enables follow-up email flows and customer segmentation inside Shopify",
    ("Already partially built in the chatbot backend. Formalising as a chargeable feature.", True),
])

add_heading("5.3  AI Kitchen Consultant Chatbot", level=2)
add_bullet([
    "Intent-based product guidance for commercial kitchen equipment",
    "Captures leads and pushes them into Shopify customer records",
    "Weekly performance reports to store owner",
    "Admin dashboard to monitor conversations and leads",
    ("Monthly subscription — covers server hosting, AI model costs, and maintenance.", True),
])

# ══════════════════════════════════════════════════════════════════════════════
# 6. PHASE 1 STATUS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Phase 1 — Current Status")
add_para(
    "All bugs within scope have been resolved. All core features are live and working. "
    "Items marked 'Needs Clarification' in the tracker require client decisions before "
    "any further development can proceed — they are not incomplete work on our side.",
    color=TEXT_MED
)
add_bullet([
    ("Theme", True),
    "Complete. All pages live. Mobile responsive. Approved design implemented.",
    ("Wholesale Pricing", True),
    "Complete across all 6 page types. CPH / ROU / non-member logic all working.",
    ("SEO & AEO", True),
    "Complete. 233 collections with meta data. Structured data live. AEO implemented.",
    ("Contact Forms", True),
    "Complete. All 4 forms working. Lead capture active. File upload pending client decision.",
    ("L2 Collection Pages", True),
    "Proper template ready. AJAX fallback live. Blocked by client not populating collections.",
    ("Chatbot", True),
    "Beta ready. Being pitched separately. Not part of Phase 1 scope.",
])

divider()

add_para(
    "Phase 1 is complete from our side. Any new requests — including file upload, chatbot "
    "subscription, or further design changes — require a new quote and client sign-off. "
    "We have delivered significantly beyond the original scope of work. "
    "This should be communicated clearly and confidently to the client.",
    color=DEEP_BLUE, bold_=True, size=11
)

divider()

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(f"Hiraya Digital  |  Prepared {datetime.date.today().strftime('%B %d, %Y')}")
r_f.font.size = Pt(9); r_f.font.color.rgb = TEXT_MED; r_f.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUTPUT_DIR, "Celebrate_Festival_Project_Summary_Disha.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
